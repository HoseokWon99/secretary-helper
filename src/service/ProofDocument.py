from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

class ProofDocument:

    def __init__(self, year: int, semester: int, proof_type: str):
        self.__doc = Document()
        self.__title = f"{year}년 {semester}학기 정기지원금 신청 - {proof_type}"
        self.__table = None
        self.__cnt = 0

    def __append_page(self):
        header = self.__doc.sections[0].header

        while len(header.paragraphs) < 2:
            header.add_paragraph()

        for i, text in enumerate([self.__title, "(건대극장)"]):
            header.paragraphs[i].text = text
            header.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header.paragraphs[i].runs[0].font.bold = True
            header.paragraphs[i].runs[0].font.size = Pt(14)

        self.__table = self.__doc.add_table(rows=4, cols=2)
        self.__table.style = 'Table Grid'
        self.__table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.__doc.add_page_break()

    def append_content(self, idx: int, image: str):
        pos = self.__cnt % 4

        if pos == 0:
            self.__append_page()

        i = 0 if pos < 2 else 2
        j = pos % 2

        p1 = self.__table.rows[i].cells[j].paragraphs[0]
        p1.text = str(idx)
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p1.runs[0].font.size = Pt(12)

        p2 = self.__table.rows[i + 1].cells[j].paragraphs[0]
        p2.add_run().add_picture(image, Inches(2))
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.__cnt += 1

    def to_docx(self, path: str):

        if not path.endswith(".docx"):
            path += ".docx"

        self.__doc.save(path)
        print(f"파일이 {path}에 성공적으로 저장되었습니다.")

